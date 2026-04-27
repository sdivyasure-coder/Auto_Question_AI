from __future__ import annotations

import csv
import hashlib
import json
import os
import re
from dataclasses import dataclass
from collections import Counter
from datetime import datetime
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
except Exception:  # pragma: no cover - optional at runtime
    canvas = None
    letter = None
    inch = None


@dataclass
class GeneratorConfig:
    difficulty_focus: int = 55
    ai_weight: int = 72
    unit_coverage: int = 80
    chapter_balance: bool = True
    question_diversity: bool = True
    strict_blueprint: bool = False

    @staticmethod
    def from_form(data: Dict[str, str]) -> "GeneratorConfig":
        def to_int(value: Optional[str], default: int) -> int:
            try:
                return max(0, min(100, int(value)))
            except (TypeError, ValueError):
                return default

        def to_bool(value: Optional[str]) -> bool:
            return str(value).lower() in ["1", "true", "on", "yes"]

        return GeneratorConfig(
            difficulty_focus=to_int(data.get("difficulty_focus"), 55),
            ai_weight=to_int(data.get("ai_weight"), 72),
            unit_coverage=to_int(data.get("unit_coverage"), 80),
            chapter_balance=to_bool(data.get("chapter_balance")),
            question_diversity=to_bool(data.get("question_diversity")),
            strict_blueprint=to_bool(data.get("strict_blueprint")),
        )


def score_question(q_text: str, mark: int, config: GeneratorConfig) -> float:
    """Score question by heuristic quality + config weighting."""
    base = mark * 10
    word_count = len(q_text.split())
    if 15 <= word_count <= 50:
        base += 20
    keywords = ["explain", "describe", "analyze", "discuss", "justify"]
    if any(kw in q_text.lower() for kw in keywords):
        base += 10

    difficulty_target = 20 + (config.difficulty_focus / 100) * 40
    difficulty_score = 20 - min(20, abs(word_count - difficulty_target))

    rand_seed = int(hashlib.md5(q_text.encode("utf-8")).hexdigest(), 16) % 100
    weighted = (config.ai_weight / 100) * (base + difficulty_score) + (
        (100 - config.ai_weight) / 100
    ) * rand_seed
    return weighted + 30


def _similarity(a: str, b: str) -> float:
    tokens_a = set(a.lower().split())
    tokens_b = set(b.lower().split())
    if not tokens_a or not tokens_b:
        return 0.0
    intersection = len(tokens_a & tokens_b)
    union = len(tokens_a | tokens_b)
    return intersection / union


def _apply_diversity(rows: pd.DataFrame, count: int, threshold: float = 0.65) -> pd.DataFrame:
    picked = []
    for _, row in rows.iterrows():
        q_text = str(row.get("question", ""))
        if not picked:
            picked.append(row)
        else:
            if all(_similarity(q_text, str(p.get("question", ""))) < threshold for p in picked):
                picked.append(row)
        if len(picked) >= count:
            break
    if not picked:
        return rows.head(count)
    return pd.DataFrame(picked)


def _round_robin_by_chapter(rows: pd.DataFrame, count: int) -> pd.DataFrame:
    if "chapter" not in rows.columns:
        return rows.head(count)
    groups = []
    for _, group in rows.groupby("chapter"):
        groups.append(group)
    picked = []
    idx = 0
    while len(picked) < count and groups:
        group = groups[idx % len(groups)]
        if not group.empty:
            picked.append(group.iloc[0])
            groups[idx % len(groups)] = group.iloc[1:]
        else:
            groups.pop(idx % len(groups))
            if not groups:
                break
        idx += 1
    if not picked:
        return rows.head(count)
    return pd.DataFrame(picked)


def select_best_questions(
    df: pd.DataFrame,
    chapter: int,
    marks: int,
    count: int,
    config: GeneratorConfig,
) -> pd.DataFrame:
    filtered = df[(df["chapter"] == chapter) & (df["mark"] == marks)].copy()
    if filtered.empty:
        return filtered
    filtered["ai_score"] = filtered["question"].apply(lambda x: score_question(str(x), marks, config))
    filtered = filtered.sort_values("ai_score", ascending=False)
    if config.question_diversity:
        threshold = 0.8 - (config.unit_coverage / 100) * 0.3
        filtered = _apply_diversity(filtered, count, threshold=threshold)
    return filtered.head(count).drop(columns=["ai_score"], errors="ignore")


def select_top_questions(
    df: pd.DataFrame,
    marks: int,
    count: int,
    config: GeneratorConfig,
) -> pd.DataFrame:
    filtered = df[df["mark"] == marks].copy()
    if filtered.empty:
        return filtered
    filtered["ai_score"] = filtered["question"].apply(lambda x: score_question(str(x), marks, config))
    filtered = filtered.sort_values("ai_score", ascending=False)
    if config.chapter_balance and config.unit_coverage >= 40:
        filtered = _round_robin_by_chapter(filtered, count)
    if config.question_diversity:
        threshold = 0.8 - (config.unit_coverage / 100) * 0.3
        filtered = _apply_diversity(filtered, count, threshold=threshold)
    return filtered.head(count).drop(columns=["ai_score"], errors="ignore")


def _warn_if_short(label: str, expected: int, actual: int, warnings: List[str]) -> None:
    if actual < expected:
        warnings.append(f"{label}: expected {expected}, found {actual}.")


def get_regular_questions(csv_file: str, subject_code: str, subject_display: str, config: GeneratorConfig) -> Dict:
    df = pd.read_csv(csv_file, engine="python", on_bad_lines="skip")
    df = df[df["subject"] == subject_code]
    warnings: List[str] = []
    data = {
        "subject": subject_display,
        "exam_title": "Regular Question Paper",
        "summary": "Pattern: 100 marks",
        "parts": {},
        "warnings": warnings,
        "config": config,
    }
    qno = 1

    qa = []
    for unit in range(1, 6):
        qs = select_best_questions(df, unit, 2, 2, config)
        _warn_if_short(f"Part A Unit {unit}", 2, len(qs), warnings)
        for _, q in qs.iterrows():
            qa.append({"no": qno, "text": q["question"], "marks": 2})
            qno += 1
    data["parts"]["A"] = {
        "title": "PART A (10 x 2 = 20 Marks)",
        "inst": "Answer ALL questions",
        "qs": qa,
    }

    qb = []
    for unit in range(1, 6):
        qs = select_best_questions(df, unit, 7, 2, config)
        _warn_if_short(f"Part B Unit {unit}", 2, len(qs), warnings)
        if len(qs) >= 2:
            qb.append({"no": qno, "a": qs.iloc[0]["question"], "b": qs.iloc[1]["question"], "marks": 7})
            qno += 1
    data["parts"]["B"] = {
        "title": "PART B (5 x 7 = 35 Marks)",
        "inst": "Answer ONE question from each pair",
        "qs": qb,
    }

    qc = []
    for unit in range(1, 6):
        qs = select_best_questions(df, unit, 15, 1, config)
        _warn_if_short(f"Part C Unit {unit}", 1, len(qs), warnings)
        for _, q in qs.iterrows():
            qc.append({"no": qno, "text": q["question"], "marks": 15})
            qno += 1
    data["parts"]["C"] = {
        "title": "PART C (3 x 15 = 45 Marks)",
        "inst": "Answer ANY THREE questions",
        "qs": qc,
    }

    if config.strict_blueprint and warnings:
        data["summary"] += " | Strict blueprint enabled: some sections are incomplete."

    return data


def get_cia_questions(csv_file: str, subject_code: str, subject_display: str, config: GeneratorConfig) -> Dict:
    df = pd.read_csv(csv_file, engine="python", on_bad_lines="skip")
    df = df[df["subject"] == subject_code]
    warnings: List[str] = []
    data = {
        "subject": subject_display,
        "exam_title": "CIA Question Paper",
        "summary": "Pattern: 50 marks (2-mark as 2, 7-mark as 6, 15-mark as 10)",
        "parts": {},
        "warnings": warnings,
        "config": config,
    }
    qno = 1

    qa = []
    qs2 = select_top_questions(df, 2, 6, config)
    _warn_if_short("Part A", 6, len(qs2), warnings)
    for _, q in qs2.iterrows():
        qa.append({"no": qno, "text": q["question"], "marks": 2})
        qno += 1
    data["parts"]["A"] = {"title": "PART A (6 x 2 = 12 Marks)", "inst": "Answer ALL questions", "qs": qa}

    qb = []
    qs_7 = select_top_questions(df, 7, 6, config).reset_index(drop=True)
    pair_count = min(len(qs_7) // 2, 3)
    _warn_if_short("Part B", 3, pair_count, warnings)
    for i in range(pair_count):
        qa_q = qs_7.iloc[i * 2]["question"]
        qb_q = qs_7.iloc[(i * 2) + 1]["question"]
        qb.append({"no": qno, "a": qa_q, "b": qb_q, "marks": 6})
        qno += 1
    data["parts"]["B"] = {
        "title": "PART B (3 x 6 = 18 Marks)",
        "inst": "Answer ONE question from each pair",
        "qs": qb,
    }

    qc = []
    qs15 = select_top_questions(df, 15, 3, config)
    _warn_if_short("Part C", 3, len(qs15), warnings)
    for _, q in qs15.iterrows():
        qc.append({"no": qno, "text": q["question"], "marks": 10})
        qno += 1
    data["parts"]["C"] = {
        "title": "PART C (3 x 10 = 30 Marks)",
        "inst": "Answer ANY TWO questions",
        "qs": qc,
    }

    if config.strict_blueprint and warnings:
        data["summary"] += " | Strict blueprint enabled: some sections are incomplete."

    return data


def get_75_questions(csv_file: str, subject_code: str, subject_display: str, config: GeneratorConfig) -> Dict:
    df = pd.read_csv(csv_file, engine="python", on_bad_lines="skip")
    df = df[df["subject"] == subject_code]
    warnings: List[str] = []
    data = {
        "subject": subject_display,
        "exam_title": "75 Mark Question Paper",
        "summary": "Pattern: 75 marks (2-mark as 2, 7-mark as 5, 15-mark as 10)",
        "parts": {},
        "warnings": warnings,
        "config": config,
    }
    qno = 1

    qa = []
    qs2 = select_top_questions(df, 2, 10, config)
    _warn_if_short("Part A", 10, len(qs2), warnings)
    for _, q in qs2.iterrows():
        qa.append({"no": qno, "text": q["question"], "marks": 2})
        qno += 1
    data["parts"]["A"] = {
        "title": "PART A (10 x 2 = 20 Marks)",
        "inst": "Answer ALL questions",
        "qs": qa,
    }

    qb = []
    qs_7 = select_top_questions(df, 7, 10, config).reset_index(drop=True)
    pair_count = min(len(qs_7) // 2, 5)
    _warn_if_short("Part B", 5, pair_count, warnings)
    for i in range(pair_count):
        qa_q = qs_7.iloc[i * 2]["question"]
        qb_q = qs_7.iloc[(i * 2) + 1]["question"]
        qb.append({"no": qno, "a": qa_q, "b": qb_q, "marks": 5})
        qno += 1
    data["parts"]["B"] = {
        "title": "PART B (5 x 5 = 25 Marks)",
        "inst": "Answer ONE question from each pair",
        "qs": qb,
    }

    qc = []
    qs15 = select_top_questions(df, 15, 3, config)
    _warn_if_short("Part C", 3, len(qs15), warnings)
    for _, q in qs15.iterrows():
        qc.append({"no": qno, "text": q["question"], "marks": 10})
        qno += 1
    data["parts"]["C"] = {
        "title": "PART C (3 x 10 = 30 Marks)",
        "inst": "Answer ALL questions",
        "qs": qc,
    }

    if config.strict_blueprint and warnings:
        data["summary"] += " | Strict blueprint enabled: some sections are incomplete."

    return data


def get_questions(csv_file: str, subject_code: str, subject_display: str, pattern: str, config: GeneratorConfig) -> Dict:
    if pattern == "cia50":
        return get_cia_questions(csv_file, subject_code, subject_display, config)
    if pattern == "75":
        return get_75_questions(csv_file, subject_code, subject_display, config)
    return get_regular_questions(csv_file, subject_code, subject_display, config)


def build_download_name(subject_display: str, pattern: str) -> str:
    subject_part = subject_display.replace(" ", "_")
    if pattern == "cia50":
        return f"{subject_part}_CIA_50_Question_Paper.docx"
    if pattern == "75":
        return f"{subject_part}_75_Mark_Question_Paper.docx"
    return f"{subject_part}_Question_Paper.docx"


def format_para(p, size: int = 12, bold: bool = False, align=None) -> None:
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    if align:
        p.alignment = align

    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold


def apply_doc_header_footer(
    doc,
    subject_title: str,
    pattern_label: str,
    max_marks: int,
    generated_by: str,
    custom_header: str = "",
    custom_footer: str = "",
) -> None:
    stamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    default_header = f"AI Question Paper Generator | {subject_title}"
    default_footer = (
        f"{pattern_label} | Max Marks: {max_marks} | "
        f"Generated: {stamp} | By: {generated_by}"
    )
    header_text = str(custom_header).strip() or default_header
    footer_text = str(custom_footer).strip() or default_footer

    for section in doc.sections:
        header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        header_para.text = header_text
        format_para(header_para, 9, False, WD_ALIGN_PARAGRAPH.CENTER)

        footer_para = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer_para.text = footer_text
        format_para(footer_para, 9, False, WD_ALIGN_PARAGRAPH.CENTER)


def apply_preview_edits(questions: Dict, edited_items: List[Dict]) -> Dict:
    if not edited_items:
        return questions
    for item in edited_items:
        try:
            part_key = item.get("part")
            index = int(item.get("index"))
            field = item.get("field")
            text = str(item.get("text", "")).strip()
            if not part_key or field not in ["text", "a", "b"]:
                continue
            if text == "":
                continue
            questions["parts"][part_key]["qs"][index][field] = text
        except (ValueError, KeyError, IndexError, TypeError):
            continue
    return questions


def generate_question_paper_from_data(
    subject_title: str,
    pattern: str,
    questions: Dict,
    generated_by: str,
    output_file: str,
    custom_header: str = "",
    custom_footer: str = "",
) -> str:
    doc = Document()

    if pattern == "cia50":
        heading = "AI GENERATED CIA QUESTION PAPER"
        time_and_marks = "Time: 2 Hours          Max Marks: 50"
        pattern_label = "CIA 50 Pattern"
        max_marks = 50
    elif pattern == "75":
        heading = "AI GENERATED 75 MARK QUESTION PAPER"
        time_and_marks = "Time: 2.5 Hours          Max Marks: 75"
        pattern_label = "75 Mark Pattern"
        max_marks = 75
    else:
        heading = "AI GENERATED QUESTION PAPER"
        time_and_marks = "Time: 3 Hours          Max Marks: 100"
        pattern_label = "Regular Pattern"
        max_marks = 100

    apply_doc_header_footer(
        doc,
        subject_title,
        pattern_label,
        max_marks,
        generated_by,
        custom_header,
        custom_footer,
    )

    p = doc.add_paragraph(heading)
    format_para(p, 14, True, WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph(
        f"Subject: {subject_title}\n"
        f"{time_and_marks}\n"
    )
    format_para(p, 12, False, WD_ALIGN_PARAGRAPH.CENTER)

    for part in questions["parts"].values():
        p = doc.add_paragraph(f"\n{part['title']}")
        format_para(p, 12, True, WD_ALIGN_PARAGRAPH.CENTER)
        p = doc.add_paragraph(part["inst"])
        format_para(p)

        for q in part["qs"]:
            if "text" in q:
                line = f"{q['no']}. {q['text']}"
                if q.get("marks"):
                    line += f" ({q['marks']} Marks)"
                p = doc.add_paragraph(line)
                format_para(p)
            else:
                p = doc.add_paragraph(f"{q['no']}. a) {q['a']}")
                format_para(p)
                p = doc.add_paragraph("    Or")
                format_para(p)
                p = doc.add_paragraph(f"    b) {q['b']}")
                format_para(p)

    doc.save(output_file)
    return output_file


def generate_pdf_from_questions(
    subject_title: str,
    pattern: str,
    questions: Dict,
    output_file: str,
    generated_by: str,
    custom_header: str = "",
    custom_footer: str = "",
) -> str:
    if canvas is None:
        raise RuntimeError("PDF generator not available. Install reportlab.")

    if pattern == "cia50":
        heading = "AI GENERATED CIA QUESTION PAPER"
        time_and_marks = "Time: 2 Hours          Max Marks: 50"
        pattern_label = "CIA 50 Pattern"
        max_marks = 50
    elif pattern == "75":
        heading = "AI GENERATED 75 MARK QUESTION PAPER"
        time_and_marks = "Time: 2.5 Hours          Max Marks: 75"
        pattern_label = "75 Mark Pattern"
        max_marks = 75
    else:
        heading = "AI GENERATED QUESTION PAPER"
        time_and_marks = "Time: 3 Hours          Max Marks: 100"
        pattern_label = "Regular Pattern"
        max_marks = 100

    stamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    header_text = custom_header.strip() or f"AI Question Paper Generator | {subject_title}"
    footer_text = custom_footer.strip() or (
        f"{pattern_label} | Max Marks: {max_marks} | Generated: {stamp} | By: {generated_by}"
    )

    c = canvas.Canvas(output_file, pagesize=letter)
    c.setTitle(f"{subject_title} Question Paper")
    width, height = letter
    y = height - 0.75 * inch

    def draw_line(text: str, font="Times-Roman", size=11, leading=16):
        nonlocal y
        if y < 1 * inch:
            c.showPage()
            y = height - 0.75 * inch
        c.setFont(font, size)
        c.drawString(0.75 * inch, y, text)
        y -= leading

    def draw_center(text: str, font="Times-Roman", size=11, leading=16):
        nonlocal y
        if y < 1 * inch:
            c.showPage()
            y = height - 0.75 * inch
        c.setFont(font, size)
        c.drawCentredString(width / 2, y, text)
        y -= leading
    
    def draw_wrapped(text: str, font="Times-Roman", size=11, leading=16, max_width=6.6 * inch):
        nonlocal y
        c.setFont(font, size)
        words = text.split()
        line = ""
        for word in words:
            test_line = f"{line} {word}".strip()
            if c.stringWidth(test_line, font, size) <= max_width:
                line = test_line
            else:
                if y < 1 * inch:
                    c.showPage()
                    y = height - 0.75 * inch
                    c.setFont(font, size)
                c.drawString(0.75 * inch, y, line)
                y -= leading
                line = word
        if line:
            if y < 1 * inch:
                c.showPage()
                y = height - 0.75 * inch
                c.setFont(font, size)
            c.drawString(0.75 * inch, y, line)
            y -= leading

    draw_center(header_text, font="Times-Bold", size=10)
    y -= 8
    draw_center(heading, font="Times-Bold", size=14, leading=20)
    draw_center(f"Subject: {subject_title}", size=12, leading=18)
    draw_center(time_and_marks, size=12, leading=18)
    y -= 10

    for part in questions["parts"].values():
        draw_center(part["title"], font="Times-Bold", size=12, leading=18)
        draw_center(part["inst"], size=11, leading=16)
        y -= 4
        for q in part["qs"]:
            if "text" in q:
                mark_text = f" ({q['marks']} Marks)" if q.get("marks") else ""
                draw_wrapped(f"{q['no']}. {q['text']}{mark_text}", size=11, leading=16)
                y -= 2
            else:
                draw_wrapped(f"{q['no']}. a) {q['a']}", size=11, leading=16)
                draw_line("    Or", size=11, leading=16)
                draw_wrapped(f"    b) {q['b']}", size=11, leading=16)
                y -= 2
        y -= 10

    c.setFont("Times-Italic", 9)
    c.drawString(0.75 * inch, 0.5 * inch, footer_text)
    c.save()
    return output_file


def get_question_bank_rows(csv_file: str) -> List[Dict]:
    if not os.path.exists(csv_file):
        return []
    df = pd.read_csv(csv_file, engine="python", on_bad_lines="skip")
    if df.empty:
        return []

    rows = []
    for idx, row in df.iterrows():
        try:
            mark = int(row.get("mark", 0))
        except (TypeError, ValueError):
            mark = 0

        if mark <= 2:
            difficulty = "Easy"
            question_type = "Short"
        elif mark <= 7:
            difficulty = "Medium"
            question_type = "Medium"
        else:
            difficulty = "Hard"
            question_type = "Long"

        rows.append(
            {
                "id": idx + 1,
                "subject": str(row.get("subject", "")),
                "chapter": row.get("chapter", ""),
                "question_text": str(row.get("question", "")),
                "mark": mark,
                "difficulty": difficulty,
                "question_type": question_type,
                "usage_freq": row.get("usage_freq", "-"),
                "source": row.get("source", "bank"),
            }
        )
    return rows


def get_question_by_id(csv_file: str, question_id: int) -> Optional[Dict]:
    rows = get_question_bank_rows(csv_file)
    for q in rows:
        if q["id"] == question_id:
            return q
    return None


def update_question_in_csv(
    csv_file: str,
    question_id: int,
    subject: str,
    chapter: int,
    mark: int,
    question_text: str,
) -> bool:
    if not os.path.exists(csv_file):
        return False
    df = pd.read_csv(csv_file, engine="python", on_bad_lines="skip")
    idx = question_id - 1
    if idx < 0 or idx >= len(df):
        return False
    df.at[idx, "subject"] = subject
    df.at[idx, "chapter"] = chapter
    df.at[idx, "mark"] = mark
    df.at[idx, "question"] = question_text
    df.to_csv(csv_file, index=False)
    return True


def append_questions_to_csv(csv_file: str, rows: List[Tuple]) -> None:
    with open(csv_file, "a", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        for row in rows:
            writer.writerow(row)


def ai_generate_questions(
    csv_file: str,
    subject: str,
    chapter: int,
    mark: int,
    count: int,
) -> List[str]:
    def _parse_llm_questions(text: str, needed: int) -> List[str]:
        cleaned: List[str] = []
        for raw in str(text or "").splitlines():
            line = raw.strip()
            if not line:
                continue
            line = re.sub(r"^\s*(?:[-*]|\d+[.)])\s*", "", line).strip()
            line = line.strip('"').strip("'").strip()
            if not line:
                continue
            if not line.endswith("?") and not line.endswith("."):
                line += "?"
            if len(line.split()) < 4:
                continue
            if line not in cleaned:
                cleaned.append(line)
            if len(cleaned) >= needed:
                break
        return cleaned[:needed]

    def _generate_with_langchain_groq(
        concept_pool: List[str],
        existing_questions: List[str],
    ) -> Optional[List[str]]:
        groq_api_key = os.getenv("GROQ_API_KEY", "").strip()
        if not groq_api_key:
            return None

        try:
            from langchain_groq import ChatGroq
        except Exception:
            return None

        model = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile").strip()
        try:
            temperature = float(os.getenv("GROQ_TEMPERATURE", "0.4"))
        except ValueError:
            temperature = 0.4

        concept_hint = ", ".join(concept_pool[:20])
        examples = "\n".join([f"- {q}" for q in existing_questions[:12]])
        prompt = (
            "You are an exam setter. Generate high-quality, non-duplicate university questions.\n"
            f"Subject code: {subject}\n"
            f"Unit/Chapter: {chapter}\n"
            f"Marks per question: {mark}\n"
            f"Number of questions required: {count}\n"
            f"Important concepts: {concept_hint}\n"
            "Avoid repeating or paraphrasing these existing questions:\n"
            f"{examples}\n"
            "Return only a JSON array of strings, no markdown, no explanation."
        )

        try:
            llm = ChatGroq(
                model=model,
                api_key=groq_api_key,
                temperature=max(0.0, min(1.0, temperature)),
            )
            response = llm.invoke(prompt)
            content = getattr(response, "content", "") or ""
        except Exception:
            return None

        try:
            parsed = json.loads(content)
            if isinstance(parsed, list):
                parsed_lines = [str(item).strip() for item in parsed if str(item).strip()]
                parsed_lines = _parse_llm_questions("\n".join(parsed_lines), count)
                if len(parsed_lines) >= count:
                    return parsed_lines[:count]
        except (TypeError, ValueError, json.JSONDecodeError):
            pass

        fallback_parse = _parse_llm_questions(str(content), count)
        if len(fallback_parse) >= count:
            return fallback_parse[:count]
        return None

    def _generate_with_heuristics(
        concept_pool: List[str],
        existing_questions: List[str],
    ) -> List[str]:
        if mark == 2:
            templates = [
                "Define {concept} in the context of {subject}.",
                "List any two key points of {concept} in Unit {chapter}.",
                "State the purpose of {concept} in {subject}.",
                "Write a short note on {concept}.",
            ]
        elif mark == 7:
            templates = [
                "Explain {concept} with suitable example(s) in {subject}.",
                "Discuss the working of {concept} and its practical significance.",
                "Analyze the advantages and limitations of {concept}.",
                "Compare {concept} with related approaches and justify their use-cases.",
            ]
        else:
            templates = [
                "Evaluate the design, challenges, and best practices of {concept} in detail.",
                "Develop a detailed discussion on {concept} with architecture/diagram and justification.",
                "Critically analyze {concept} with performance, security, and scalability considerations.",
                "Propose a solution approach using {concept} for a realistic scenario in {subject}.",
            ]

        generated: List[str] = []
        max_attempts = max(count * 8, 20)
        attempts = 0
        while len(generated) < count and attempts < max_attempts:
            concept = concept_pool[(attempts * 3) % len(concept_pool)]
            template = templates[(attempts * 5) % len(templates)]
            q = template.format(concept=concept, subject=subject, chapter=chapter)
            q = q[0].upper() + q[1:] if q else q

            duplicate_generated = any(_similarity(q, g) >= 0.72 for g in generated)
            duplicate_existing = any(_similarity(q, e) >= 0.82 for e in existing_questions[:200])
            if not duplicate_generated and not duplicate_existing:
                generated.append(q)
            attempts += 1

        idx = 0
        while len(generated) < count:
            concept = concept_pool[idx % len(concept_pool)]
            generated.append(f"Explain the role of {concept} in {subject} (Unit {chapter}).")
            idx += 1
        return generated[:count]

    df = pd.read_csv(csv_file, engine="python", on_bad_lines="skip")
    pool = df[(df["subject"] == subject) & (df["chapter"] == chapter)].copy()
    if pool.empty:
        pool = df[df["subject"] == subject].copy()

    # Lightweight concept extraction from existing bank questions.
    stop_words = {
        "the", "and", "for", "with", "from", "into", "that", "this", "what", "when", "where",
        "which", "why", "how", "list", "state", "mention", "explain", "describe", "discuss",
        "analyze", "evaluate", "compare", "unit", "chapter", "marks", "question", "questions",
        "about", "your", "their", "using", "used", "between", "among", "of", "in", "to", "on",
        "is", "are", "be", "by", "an", "a",
        "suitable", "diagram", "detail", "points", "purpose", "short", "note", "write",
    }
    generic_terms = {
        "suitable", "diagram", "detail", "points", "purpose", "short", "note",
        "working", "approach", "method", "methods", "system", "systems",
    }

    word_counter: Counter = Counter()
    phrase_counter: Counter = Counter()
    existing_questions = [str(q).strip() for q in pool.get("question", pd.Series(dtype=str)).astype(str).tolist()]

    for text in existing_questions:
        tokens = [t for t in re.findall(r"[a-zA-Z][a-zA-Z0-9\-]{2,}", text.lower()) if t not in stop_words]
        word_counter.update(tokens)
        for i in range(len(tokens) - 1):
            a, b = tokens[i], tokens[i + 1]
            if a not in stop_words and b not in stop_words:
                phrase_counter.update([f"{a} {b}"])

    extracted_concepts: List[str] = [p for p, _ in phrase_counter.most_common(24)]
    extracted_concepts.extend([w for w, _ in word_counter.most_common(24)])

    # Subject-aware fallback concepts when bank data is sparse/noisy.
    defaults_by_subject = {
        "CNCC": ["network architecture", "routing protocols", "transport layer", "congestion control", "cloud deployment models"],
        "CNS": ["network security", "firewall rules", "encryption methods", "threat modeling", "access control"],
        "MA": ["android lifecycle", "ui components", "data persistence", "app architecture", "intent communication"],
    }
    fallback_concepts = defaults_by_subject.get(subject, ["core concepts", "system design", "performance optimization"])
    concepts: List[str] = list(fallback_concepts)
    for c in extracted_concepts:
        if c not in concepts:
            concepts.append(c)

    # Deduplicate while preserving order.
    seen = set()
    concept_pool = []
    for c in concepts:
        cc = c.strip()
        cc_tokens = [t for t in cc.split() if t]
        informative = any(t not in generic_terms for t in cc_tokens)
        if cc and informative and cc not in seen:
            seen.add(cc)
            concept_pool.append(cc)
    if not concept_pool:
        concept_pool = fallback_concepts
    generated = _generate_with_langchain_groq(concept_pool, existing_questions)
    if generated:
        return generated[:count]
    return _generate_with_heuristics(concept_pool, existing_questions)
