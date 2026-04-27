import csv
import json
import os
import random
import re
from collections import Counter, defaultdict
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from sqlalchemy.orm import Session

from app.models.app_setting import AppSetting
from app.models.generated_paper import GeneratedPaper, GeneratedPaperQuestion
from app.models.paper_template import PaperTemplate
from app.models.question import Question
from app.models.question_profile import QuestionProfile
from app.models.subject import Subject
from app.models.unit import Unit
from app.models.upload_asset import UploadAsset

OUTPUT_DIR = "output/papers_v2"
UPLOAD_DIR = "output/uploads"


def _ensure_dirs() -> None:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(UPLOAD_DIR, exist_ok=True)


def _json_load(raw: Optional[str], default):
    if not raw:
        return default
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return default


def seed_default_templates(db: Session) -> None:
    defaults = [
        {
            "name": "internal",
            "template_type": "internal",
            "total_marks": 50,
            "structure": {
                "sections": [
                    {"name": "A", "mark": 2, "count": 6},
                    {"name": "B", "mark": 7, "count": 3},
                    {"name": "C", "mark": 15, "count": 2},
                ]
            },
        },
        {
            "name": "model",
            "template_type": "model",
            "total_marks": 75,
            "structure": {
                "sections": [
                    {"name": "A", "mark": 2, "count": 10},
                    {"name": "B", "mark": 7, "count": 5},
                    {"name": "C", "mark": 15, "count": 3},
                ]
            },
        },
        {
            "name": "semester",
            "template_type": "semester",
            "total_marks": 100,
            "structure": {
                "sections": [
                    {"name": "A", "mark": 2, "count": 10},
                    {"name": "B", "mark": 7, "count": 5},
                    {"name": "C", "mark": 15, "count": 3},
                ]
            },
        },
    ]

    for item in defaults:
        existing = db.query(PaperTemplate).filter(PaperTemplate.name == item["name"]).first()
        if existing:
            continue
        db.add(
            PaperTemplate(
                name=item["name"],
                template_type=item["template_type"],
                total_marks=item["total_marks"],
                structure_json=json.dumps(item["structure"]),
                is_active=True,
            )
        )
    db.commit()


def _difficulty_for_mark(mark: int) -> str:
    if mark <= 2:
        return "easy"
    if mark <= 7:
        return "medium"
    return "hard"


def _blooms_from_text(text: str) -> str:
    t = text.lower()
    if any(k in t for k in ["define", "list", "state", "name"]):
        return "remember"
    if any(k in t for k in ["explain", "describe", "summarize"]):
        return "understand"
    if any(k in t for k in ["apply", "solve", "implement", "demonstrate"]):
        return "apply"
    return "analyze"


def _default_section_plan(total_marks: int, question_count: int) -> List[Dict[str, int]]:
    section_marks = [("A", 2, 0.30), ("B", 7, 0.35), ("C", 15, 0.35)]
    counts: List[Tuple[str, int, int]] = []
    for name, mark, ratio in section_marks:
        target = max(1, round((total_marks * ratio) / mark))
        counts.append((name, mark, target))

    if question_count > 0:
        current = sum(c for _, _, c in counts)
        if current > 0:
            scale = question_count / current
            scaled = []
            for name, mark, count in counts:
                scaled.append((name, mark, max(1, round(count * scale))))
            counts = scaled

    return [{"section": n, "mark": m, "count": c} for n, m, c in counts]


def _extract_concepts(rows: List[Question]) -> List[str]:
    stop = {
        "the",
        "and",
        "for",
        "with",
        "from",
        "that",
        "this",
        "what",
        "which",
        "where",
        "when",
        "is",
        "are",
        "a",
        "an",
        "in",
        "of",
        "to",
    }
    tokens: Counter = Counter()
    for row in rows:
        words = re.findall(r"[a-zA-Z][a-zA-Z0-9\-]{2,}", (row.text or "").lower())
        tokens.update([w for w in words if w not in stop])
    concepts = [w for w, _ in tokens.most_common(20)]
    if not concepts:
        return ["core concept", "architecture", "workflow", "security", "performance"]
    return concepts


def _ai_generate_question(subject_name: str, unit_no: Optional[int], mark: int, difficulty: str, concept: str) -> str:
    unit_text = f"Unit {unit_no}" if unit_no else "this subject"
    if difficulty == "easy":
        return f"Define {concept} and mention its significance in {subject_name} ({unit_text})."
    if difficulty == "medium":
        return f"Explain the working of {concept} with an example in {subject_name} ({unit_text})."
    return (
        f"Analyze a problem-solving scenario using {concept} in {subject_name} ({unit_text}) "
        "and justify the solution approach."
    )


def _map_co_po(unit_no: Optional[int], blooms_level: str) -> Tuple[str, str]:
    co = f"CO{max(1, (unit_no or 1))}"
    po_map = {"remember": "PO1", "understand": "PO2", "apply": "PO3", "analyze": "PO4"}
    return co, po_map.get(blooms_level, "PO2")


def _build_preview_html(subject: Subject, paper: GeneratedPaper, items: List[GeneratedPaperQuestion]) -> str:
    rows = []
    for item in sorted(items, key=lambda x: (x.section, x.question_no)):
        rows.append(
            f"<p><strong>{item.question_no}.</strong> {item.text} <span style='float:right;'>[{item.marks}]</span></p>"
        )
    return (
        "<div class='paper-sheet'>"
        f"<h2 style='text-align:center;margin:0;'>{subject.name} - Question Paper</h2>"
        f"<p style='text-align:center;margin:6px 0;'>Max Marks: {paper.total_marks}</p>"
        "<hr/>"
        + "".join(rows)
        + "</div>"
    )


def _generate_answer_key(items: List[GeneratedPaperQuestion]) -> Dict[str, str]:
    out: Dict[str, str] = {}
    for item in items:
        if item.marks <= 2:
            answer = "Provide concise definition and two key points."
        elif item.marks <= 7:
            answer = "Explain concept, workflow, and one practical example."
        else:
            answer = "Present structured solution with analysis, assumptions, and justification."
        out[str(item.question_no)] = answer
        item.answer_key = answer
    return out


def _quality_report(items: List[GeneratedPaperQuestion], requested_difficulty: str) -> Dict[str, Any]:
    if not items:
        return {"status": "empty", "coverage": {}, "difficulty_balance": {}, "repetition_risk": 0}

    by_diff: Counter = Counter([i.difficulty for i in items])
    by_section: Counter = Counter([i.section for i in items])

    repetition = 0
    seen = set()
    for item in items:
        key = re.sub(r"\s+", " ", item.text.strip().lower())
        if key in seen:
            repetition += 1
        seen.add(key)

    return {
        "status": "ok",
        "topic_coverage": dict(by_section),
        "difficulty_balance": dict(by_diff),
        "repetition_risk": repetition,
        "requested_difficulty": requested_difficulty,
    }


def _estimate_duration(items: List[GeneratedPaperQuestion]) -> int:
    return int(sum(max(2, item.marks) for item in items) * 1.8)


def _select_candidates(
    db: Session,
    subject_id: int,
    allowed_units: List[int],
    avoid_question_ids: set,
    mark: int,
    count: int,
    requested_difficulty: str,
) -> List[Question]:
    query = db.query(Question).filter(
        Question.subject_id == subject_id,
        Question.mark == mark,
        Question.active.is_(True),
    )
    if allowed_units:
        query = query.filter(Question.chapter.in_(allowed_units))

    rows = query.all()
    random.shuffle(rows)

    selected: List[Question] = []
    for row in rows:
        if row.id in avoid_question_ids:
            continue
        if requested_difficulty in ["easy", "medium", "hard"] and row.difficulty != requested_difficulty:
            continue
        selected.append(row)
        avoid_question_ids.add(row.id)
        if len(selected) >= count:
            return selected

    # fallback ignoring difficulty
    for row in rows:
        if row.id in avoid_question_ids:
            continue
        selected.append(row)
        avoid_question_ids.add(row.id)
        if len(selected) >= count:
            break
    return selected


def _previous_question_ids(db: Session, subject_id: int) -> set:
    ids = set()
    recent = db.query(GeneratedPaper).filter(GeneratedPaper.subject_id == subject_id).order_by(GeneratedPaper.id.desc()).limit(25).all()
    for paper in recent:
        for item in paper.questions:
            if item.question_id:
                ids.add(item.question_id)
    return ids


def _get_template_plan(db: Session, template_name: str, total_marks: int, question_count: int) -> List[Dict[str, int]]:
    template = db.query(PaperTemplate).filter(PaperTemplate.name == template_name).first()
    if template:
        structure = _json_load(template.structure_json, {})
        sections = structure.get("sections") or []
        if sections:
            return [{"section": s["name"], "mark": int(s["mark"]), "count": int(s["count"])} for s in sections]
    return _default_section_plan(total_marks, question_count)


def generate_paper_v2(
    db: Session,
    user_id: int,
    subject_id: int,
    unit_ids: List[int],
    total_marks: int,
    difficulty: str,
    question_count: int,
    template_name: str,
    exam_date: Optional[str] = None,
) -> GeneratedPaper:
    _ensure_dirs()
    seed_default_templates(db)

    subject = db.query(Subject).filter(Subject.id == subject_id).first()
    if not subject:
        raise ValueError("Subject not found")

    selected_units = db.query(Unit).filter(Unit.id.in_(unit_ids)).all() if unit_ids else []
    allowed_unit_numbers = [u.unit_no for u in selected_units]

    template = db.query(PaperTemplate).filter(PaperTemplate.name == template_name).first()
    section_plan = _get_template_plan(db, template_name, total_marks, question_count)

    paper = GeneratedPaper(
        subject_id=subject_id,
        user_id=user_id,
        template_id=template.id if template else None,
        title=f"{subject.code} {template_name.title()} Paper {datetime.now().strftime('%Y%m%d_%H%M%S')}",
        total_marks=total_marks,
        requested_difficulty=difficulty,
        num_questions=0,
        selected_units_json=json.dumps(allowed_unit_numbers),
        status="draft",
        exam_date=exam_date,
    )
    db.add(paper)
    db.commit()
    db.refresh(paper)

    avoid_ids = _previous_question_ids(db, subject_id)
    generated_items: List[GeneratedPaperQuestion] = []
    serial = 1

    # collect subject pool once for AI concept fallback
    pool = db.query(Question).filter(Question.subject_id == subject_id, Question.active.is_(True)).all()
    concepts = _extract_concepts(pool)
    cidx = 0

    for sec in section_plan:
        section = sec["section"]
        mark = sec["mark"]
        count = sec["count"]
        picked = _select_candidates(
            db,
            subject_id,
            allowed_unit_numbers,
            avoid_ids,
            mark,
            count,
            difficulty.lower(),
        )

        for q in picked:
            profile = db.query(QuestionProfile).filter(QuestionProfile.question_id == q.id).first()
            blooms = profile.blooms_level if profile else _blooms_from_text(q.text)
            co, po = _map_co_po(q.chapter, blooms)
            generated_items.append(
                GeneratedPaperQuestion(
                    generated_paper_id=paper.id,
                    question_id=q.id,
                    section=section,
                    question_no=serial,
                    marks=mark,
                    text=q.text,
                    difficulty=q.difficulty or _difficulty_for_mark(mark),
                    blooms_level=blooms,
                    co_code=profile.co_code if profile and profile.co_code else co,
                    po_code=profile.po_code if profile and profile.po_code else po,
                    is_ai_generated=0,
                )
            )
            q.usage_count = int(q.usage_count or 0) + 1
            serial += 1

        missing = max(0, count - len(picked))
        for _ in range(missing):
            concept = concepts[cidx % len(concepts)]
            cidx += 1
            unit_no = random.choice(allowed_unit_numbers) if allowed_unit_numbers else None
            ai_q = _ai_generate_question(subject.name, unit_no, mark, difficulty.lower(), concept)
            blooms = _blooms_from_text(ai_q)
            co, po = _map_co_po(unit_no, blooms)
            generated_items.append(
                GeneratedPaperQuestion(
                    generated_paper_id=paper.id,
                    question_id=None,
                    section=section,
                    question_no=serial,
                    marks=mark,
                    text=ai_q,
                    difficulty=difficulty.lower() if difficulty.lower() in ["easy", "medium", "hard"] else _difficulty_for_mark(mark),
                    blooms_level=blooms,
                    co_code=co,
                    po_code=po,
                    is_ai_generated=1,
                )
            )
            serial += 1

    random.shuffle(generated_items)
    for i, item in enumerate(generated_items, start=1):
        item.question_no = i

    db.add_all(generated_items)
    paper.questions = generated_items

    answers = _generate_answer_key(generated_items)
    quality = _quality_report(generated_items, difficulty.lower())
    paper.answer_key_json = json.dumps(answers)
    paper.quality_report_json = json.dumps(quality)
    paper.num_questions = len(generated_items)
    paper.estimated_minutes = _estimate_duration(generated_items)
    paper.preview_html = _build_preview_html(subject, paper, generated_items)
    db.commit()
    db.refresh(paper)

    export_paper_files(db, paper.id)
    return db.query(GeneratedPaper).filter(GeneratedPaper.id == paper.id).first()


def export_paper_files(db: Session, paper_id: int) -> GeneratedPaper:
    _ensure_dirs()
    paper = db.query(GeneratedPaper).filter(GeneratedPaper.id == paper_id).first()
    if not paper:
        raise ValueError("Paper not found")
    subject = db.query(Subject).filter(Subject.id == paper.subject_id).first()
    if not subject:
        raise ValueError("Subject not found")

    items = sorted(paper.questions, key=lambda x: x.question_no)

    pdf_path = os.path.join(OUTPUT_DIR, f"generated_{paper.id}.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y = height - 0.8 * inch

    def draw_line(text: str, size: int = 11, bold: bool = False):
        nonlocal y
        if y < 1.0 * inch:
            c.showPage()
            y = height - 0.8 * inch
        c.setFont("Times-Bold" if bold else "Times-Roman", size)
        c.drawString(0.75 * inch, y, text)
        y -= 15

    draw_line(f"{subject.name} - Question Paper", size=13, bold=True)
    draw_line(f"Total Marks: {paper.total_marks} | Date: {paper.exam_date or '-'}", size=10)
    y -= 4
    for item in items:
        draw_line(f"{item.question_no}. {item.text} ({item.marks} marks)", size=11)

    c.save()

    docx_path = os.path.join(OUTPUT_DIR, f"generated_{paper.id}.docx")
    doc = Document()
    doc.add_heading(f"{subject.name} - Question Paper", level=1)
    meta = doc.add_paragraph(f"Total Marks: {paper.total_marks}   Date: {paper.exam_date or '-'}")
    for run in meta.runs:
        run.font.size = Pt(10)
    for item in items:
        doc.add_paragraph(f"{item.question_no}. {item.text} ({item.marks} marks)")
    doc.save(docx_path)

    paper.file_pdf = pdf_path
    paper.file_docx = docx_path
    paper.status = "final"
    db.commit()
    db.refresh(paper)
    return paper


def save_draft(db: Session, paper_id: int) -> GeneratedPaper:
    paper = db.query(GeneratedPaper).filter(GeneratedPaper.id == paper_id).first()
    if not paper:
        raise ValueError("Paper not found")
    paper.status = "draft"
    db.commit()
    db.refresh(paper)
    return paper


def regenerate_section(db: Session, paper_id: int, section: str) -> GeneratedPaper:
    paper = db.query(GeneratedPaper).filter(GeneratedPaper.id == paper_id).first()
    if not paper:
        raise ValueError("Paper not found")
    subject = db.query(Subject).filter(Subject.id == paper.subject_id).first()
    if not subject:
        raise ValueError("Subject not found")

    items = [q for q in paper.questions if q.section.upper() == section.upper()]
    if not items:
        return paper

    pool = db.query(Question).filter(Question.subject_id == paper.subject_id, Question.active.is_(True)).all()
    concepts = _extract_concepts(pool)
    for idx, item in enumerate(items):
        concept = concepts[(idx * 2) % len(concepts)]
        item.text = _ai_generate_question(subject.name, None, item.marks, item.difficulty, concept)
        item.is_ai_generated = 1
        item.blooms_level = _blooms_from_text(item.text)

    answers = _generate_answer_key(paper.questions)
    paper.answer_key_json = json.dumps(answers)
    paper.quality_report_json = json.dumps(_quality_report(paper.questions, paper.requested_difficulty))
    paper.preview_html = _build_preview_html(subject, paper, paper.questions)
    db.commit()
    db.refresh(paper)
    return paper


def shuffle_questions(db: Session, paper_id: int) -> GeneratedPaper:
    paper = db.query(GeneratedPaper).filter(GeneratedPaper.id == paper_id).first()
    if not paper:
        raise ValueError("Paper not found")
    items = list(paper.questions)
    random.shuffle(items)
    for idx, item in enumerate(items, start=1):
        item.question_no = idx
    paper.preview_html = _build_preview_html(paper.subject, paper, items)
    db.commit()
    db.refresh(paper)
    return paper


def improve_question_text(question_text: str) -> str:
    text = re.sub(r"\s+", " ", (question_text or "").strip())
    if not text:
        return ""
    starters = ["Explain", "Discuss", "Analyze", "Evaluate", "Illustrate"]
    if not any(text.lower().startswith(s.lower()) for s in starters):
        text = f"Explain {text[0].lower() + text[1:]}"
    if not text.endswith("?"):
        text = text.rstrip(".") + "."
    return text


def save_setting(db: Session, key: str, value: Dict[str, Any]) -> AppSetting:
    setting = db.query(AppSetting).filter(AppSetting.key == key).first()
    raw = json.dumps(value)
    if setting:
        setting.value_json = raw
    else:
        setting = AppSetting(key=key, value_json=raw)
        db.add(setting)
    db.commit()
    db.refresh(setting)
    return setting


def get_setting(db: Session, key: str, default: Dict[str, Any]) -> Dict[str, Any]:
    setting = db.query(AppSetting).filter(AppSetting.key == key).first()
    if not setting:
        return default
    return _json_load(setting.value_json, default)


def save_upload_file(db: Session, user_id: int, subject_id: Optional[int], file_type: str, filename: str, content: bytes) -> UploadAsset:
    _ensure_dirs()
    safe_name = re.sub(r"[^a-zA-Z0-9_.-]", "_", filename)
    path = os.path.join(UPLOAD_DIR, f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{safe_name}")
    with open(path, "wb") as f:
        f.write(content)
    upload = UploadAsset(
        subject_id=subject_id,
        user_id=user_id,
        file_type=file_type,
        file_path=path,
        parse_status="pending",
    )
    db.add(upload)
    db.commit()
    db.refresh(upload)
    return upload


def extract_questions_from_upload(db: Session, upload_id: int, subject_id: int) -> Tuple[int, str]:
    upload = db.query(UploadAsset).filter(UploadAsset.id == upload_id).first()
    if not upload:
        raise ValueError("Upload not found")

    inserted = 0
    lines: List[str] = []

    try:
        with open(upload.file_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = [line.strip() for line in f.readlines() if line.strip()]
    except OSError:
        upload.parse_status = "failed"
        upload.notes = "Unable to read file"
        db.commit()
        return 0, "failed"

    candidates = []
    for line in lines:
        if "?" in line or len(line.split()) > 7:
            candidates.append(line)

    for line in candidates:
        text = line.strip()
        if not text:
            continue
        mark = 2
        l = text.lower()
        if any(k in l for k in ["analyze", "evaluate", "justify", "case study"]):
            mark = 15
        elif any(k in l for k in ["explain", "discuss", "illustrate", "compare"]):
            mark = 7

        chapter_match = re.search(r"unit\s*(\d+)|chapter\s*(\d+)", l)
        chapter = 1
        if chapter_match:
            chapter = int(chapter_match.group(1) or chapter_match.group(2))

        q = Question(
            subject_id=subject_id,
            chapter=chapter,
            text=text,
            mark=mark,
            difficulty=_difficulty_for_mark(mark),
            source="prev_year" if upload.file_type == "prev_year" else "syllabus",
            active=True,
        )
        db.add(q)
        db.flush()

        profile = QuestionProfile(
            question_id=q.id,
            blooms_level=_blooms_from_text(text),
            co_code=f"CO{chapter}",
            po_code="PO2",
        )
        db.add(profile)
        inserted += 1

    upload.parse_status = "done"
    upload.notes = f"Inserted {inserted} questions"
    db.commit()
    return inserted, "done"


def history_search(
    db: Session,
    user_id: int,
    subject_id: Optional[int] = None,
    query: str = "",
) -> List[GeneratedPaper]:
    q = db.query(GeneratedPaper).filter(GeneratedPaper.user_id == user_id)
    if subject_id:
        q = q.filter(GeneratedPaper.subject_id == subject_id)
    if query:
        q = q.filter(GeneratedPaper.title.ilike(f"%{query}%"))
    return q.order_by(GeneratedPaper.id.desc()).all()


def unit_weightage(db: Session, paper_id: int) -> Dict[str, Any]:
    paper = db.query(GeneratedPaper).filter(GeneratedPaper.id == paper_id).first()
    if not paper:
        raise ValueError("Paper not found")

    unit_marks: Dict[int, int] = defaultdict(int)
    total = 0
    for item in paper.questions:
        total += item.marks
        if item.question_id:
            q = db.query(Question).filter(Question.id == item.question_id).first()
            if q:
                unit_marks[int(q.chapter)] += int(item.marks)

    out = [{"unit": k, "marks": v} for k, v in sorted(unit_marks.items())]
    return {"paper_id": paper_id, "unit_marks": out, "total_marks": total}


def analytics_dashboard(db: Session) -> Dict[str, Any]:
    papers = db.query(GeneratedPaper).all()
    if not papers:
        return {"most_used_units": [], "generation_frequency": [], "difficulty_distribution": []}

    unit_counter: Counter = Counter()
    date_counter: Counter = Counter()
    diff_counter: Counter = Counter()

    for paper in papers:
        date_counter[paper.created_at.strftime("%Y-%m-%d") if paper.created_at else "unknown"] += 1
        diff_counter[paper.requested_difficulty] += 1
        for item in paper.questions:
            if item.question_id:
                q = db.query(Question).filter(Question.id == item.question_id).first()
                if q:
                    unit_counter[f"Unit {q.chapter}"] += 1

    most_used_units = [{"unit": k, "count": v} for k, v in unit_counter.most_common(10)]
    generation_frequency = [{"date": k, "count": v} for k, v in sorted(date_counter.items())]
    difficulty_distribution = [{"difficulty": k, "count": v} for k, v in diff_counter.items()]
    return {
        "most_used_units": most_used_units,
        "generation_frequency": generation_frequency,
        "difficulty_distribution": difficulty_distribution,
    }


def paper_to_dict(paper: GeneratedPaper) -> Dict[str, Any]:
    quality = _json_load(paper.quality_report_json, {})
    answers = _json_load(paper.answer_key_json, {})
    questions = sorted(paper.questions, key=lambda x: x.question_no)

    pdf_url = None
    docx_url = None
    if paper.file_pdf:
        rel = paper.file_pdf.replace("output\\", "").replace("output/", "").replace("\\", "/")
        pdf_url = f"/files/{rel}"
    if paper.file_docx:
        rel = paper.file_docx.replace("output\\", "").replace("output/", "").replace("\\", "/")
        docx_url = f"/files/{rel}"

    return {
        "id": paper.id,
        "subject_id": paper.subject_id,
        "title": paper.title,
        "total_marks": paper.total_marks,
        "requested_difficulty": paper.requested_difficulty,
        "num_questions": paper.num_questions,
        "status": paper.status,
        "exam_date": paper.exam_date,
        "estimated_minutes": paper.estimated_minutes,
        "quality_report": quality,
        "answer_keys": answers,
        "preview_html": paper.preview_html,
        "pdf_url": pdf_url,
        "docx_url": docx_url,
        "questions": [
            {
                "id": q.id,
                "section": q.section,
                "question_no": q.question_no,
                "marks": q.marks,
                "text": q.text,
                "difficulty": q.difficulty,
                "blooms_level": q.blooms_level,
                "co_code": q.co_code,
                "po_code": q.po_code,
                "answer_key": q.answer_key,
            }
            for q in questions
        ],
    }
